"""
interview_manager.py

This module defines the InterviewManager class, which manages the interview
process, including handling questions, transcribing audio, and generating
transcripts. It interacts with the OpenAI API to facilitate the conversation
and evaluate candidate responses.
"""

import os
import shutil
import subprocess
from datetime import datetime
from functools import cached_property
from string import Template

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from commons import DBEnterExitMixin, FileName, InterviewStatus, logger
from interview.src.models import InterviewORM
from interview.src.prompt import INSTRUCTIONS
from user_management.src.candidate import Candidate

load_dotenv(override=True)
FILES_DIR = os.environ.get("FILES_DIR")


class InterviewManager(DBEnterExitMixin):
    """Manage the interview process, including handling questions,
    transcribing audio, and generating transcripts.

    Attributes:
        candidate_id (int): The ID of the candidate being interviewed.
    """

    def __init__(self, candidate_id: int):
        self.output_dir = os.path.join(FILES_DIR, str(candidate_id))
        os.makedirs(self.output_dir, exist_ok=True)
        self.__candidate = Candidate(user_id=candidate_id)
        self.qa_dir = os.path.join(
            FILES_DIR,
            self.__candidate.user_profile.skill,
            self.__candidate.user_profile.designation,
        )
        self._db_helper = None

    def stop(self, deactivate=True):
        """Stop the interview process.

        This method finalizes the interview and performs any necessary cleanup.
        """

        if deactivate:
            with self:
                self._db_session.query(InterviewORM).filter(
                    InterviewORM.candidate_id == self.__candidate.user_id
                ).filter(
                    InterviewORM.status == InterviewStatus.IN_PROGRESS.value
                ).update(
                    {"status": InterviewStatus.COMPLETED.value}
                )
                self._db_session.commit()
            self.__candidate.deactivate()

    def move_status_to_in_progress(self):
        """
        Update the interview status of the candidate.

        This method updates the interview status of the candidate identified by
        the current user ID to the given status. The status is updated in the
        database and the method returns after committing the changes.

        Args:
            status (InterviewStatus): The new status of the candidate.
        """
        with self:
            self._db_session.query(InterviewORM).filter(
                InterviewORM.candidate_id == self.__candidate.user_id
            ).filter(InterviewORM.status == InterviewStatus.SCHEDULED.value).update(
                {
                    "status": InterviewStatus.IN_PROGRESS.value,
                    "interview_datetime": datetime.now(),
                }
            )
            self._db_session.commit()

    # def __del__(self):
    #     if not self.__stop_called:
    #         self.stop()

    # from moviepy import VideoFileClip
    # def __extract_audio(self):
    #     clip = VideoFileClip(self.__candidate_speech_path + "mp4")
    #     clip.audio.write_audiofile(self.__candidate_speech_path + "wav")
    #     clip.close()

    def generate_transcript(self, transcript: list):
        """
        Generates a PDF transcript of the interview from the given transcript
        text list.

        Args:
            transcript (list): A list of strings, where each string is a line
                of text from the transcript.

        This method generates a PDF transcript of the interview and saves it to
        the `output_dir` path specified in the `__init__` method.
        """
        file_name = f"{FileName.TRANSCRIPT.value}"
        logger.info("Generating %s.txt ..", file_name)
        with open(
            os.path.join(self.output_dir, f"{file_name}.txt"), "w", encoding="utf-8"
        ) as f:
            f.write("\n".join(transcript))

        logger.info("Generating transcript PDF...")
        # Define PDF file path
        pdf_path = os.path.join(self.output_dir, f"{file_name}.pdf")

        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)

        # Get default styles
        styles = getSampleStyleSheet()
        style = styles["BodyText"]  # Ensures text wraps automatically

        # Create a list of Paragraph objects
        content = []
        for entry in transcript:
            content.append(Paragraph(entry, style))
            content.append(Spacer(1, 10))  # Add spacing between lines

        # Build the PDF document
        doc.build(content)
        logger.debug("Transcript PDF generated and saved to %s", pdf_path)

    def get_instructions(self):
        """
        Generates the interview instructions based on the candidate's skill,
        job description, resume, and important questions.

        Returns:
            str: The generated interview instructions.
        """
        instructions = Template(INSTRUCTIONS)
        jd = self.__get_jd()
        important_questions = self.__get_question_bank()
        return instructions.substitute(
            skill_set=self.__candidate.user_profile.skill,
            job_description=jd,
            candidate_resume=self.__resume,
            important_questions=important_questions,
            candidate_name=self.__candidate.user_profile.name,
        )
        # print(instructions)

    def __get_question_bank(self):
        file_path = os.path.join(self.qa_dir, f"{FileName.QUESTION_BANK.value}.docx")
        try:
            doc = Document(file_path)
        except (FileNotFoundError, PackageNotFoundError) as e:
            logger.error(
                "Question bank not found for skill: %s and designation: %s",
                self.__candidate.user_profile.skill,
                self.__candidate.user_profile.designation,
            )
            logger.exception(e)
            return ""
        questions = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Check if the paragraph is not empty
                questions.append(f"{i + 1}. {para.text}")  # Add numbering
        return "\n".join(questions)

    def pre_requisites(self):
        """
        Checks if all the required files for the interview are present.

        Returns:
            bool: True if all the required files are present, False otherwise.
        """
        return all([self.__get_jd(), self.__resume, self.__get_question_bank()])

    def __get_jd(self):
        file_path = os.path.join(self.qa_dir, f"{FileName.JD.value}.docx")
        try:
            doc = Document(file_path)
            jd = ""
            for para in doc.paragraphs:
                jd += para.text + "\n"
            return jd
        except PackageNotFoundError as e:
            logger.error(
                "Job description not found for skill: %s and designation: %s",
                self.__candidate.user_profile.skill,
                self.__candidate.user_profile.designation,
            )
            logger.exception(e)
            return ""

    @cached_property
    def __resume(self) -> str:
        docx_path = os.path.join(self.output_dir, f"{FileName.RESUME.value}.docx")
        pdf_path = os.path.join(self.output_dir, f"{FileName.RESUME.value}.pdf")

        # Check if either file exists
        if not (os.path.exists(docx_path) or os.path.exists(pdf_path)):
            logger.warning(
                "No resume file found for candidate id: %s",
                self.__candidate.user_profile.id,
            )
            return ""

        try:
            if os.path.exists(docx_path):
                logger.debug("Reading %s", docx_path)
                doc = Document(docx_path)
                return "\n".join([para.text for para in doc.paragraphs])

            if os.path.exists(pdf_path):
                logger.debug("Reading %s", pdf_path)
                with open(pdf_path, "rb") as f:
                    pdf_reader = PdfReader(f)
                    return "\n".join([page.extract_text() for page in pdf_reader.pages])

        except (PackageNotFoundError, PdfReadError) as e:
            logger.error(
                "Failed to read resume for candidate id: %s",
                self.__candidate.user_profile.id,
            )
            logger.exception(e)
        return ""

    def remux_video(self) -> str:
        """
        Remuxes a video file (WebM/MP4) to make it seekable by fixing metadata placement.
        Uses ffmpeg with -c copy for fast, lossless remuxing.

        Args:
            input_path (str): Path to the original video file.
            output_path (str, optional): Path for the remuxed file. If None, appends '_remuxed' to input filename.

        Returns:
            str: Path to the remuxed video file.

        Raises:
            RuntimeError: If ffmpeg fails.
        """
        video_file_path = os.path.join(self.output_dir, f"{FileName.VIDEO.value}.webm")
        output_path = os.path.join(
            self.output_dir, f"{FileName.VIDEO.value}_remuxed.webm"
        )
        cmd = [
            "ffmpeg",
            "-y",  # Overwrite output file if exists
            "-i",
            video_file_path,
            "-c",
            "copy",  # Copy streams (no re-encoding)
            output_path,
        ]

        try:
            subprocess.run(
                cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
            )
            # Only overwrite if remuxed file exists and is non-empty
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                shutil.move(output_path, video_file_path)
                logger.info(
                    "Remuxed video saved and replaced original: %s", video_file_path
                )
            else:
                logger.warning(
                    "Remuxed file not created or empty, keeping original: %s",
                    video_file_path,
                )
        except subprocess.CalledProcessError as e:
            logger.exception("ffmpeg remux failed: %s", e.stderr.decode("utf-8"))
            logger.exception(e)
